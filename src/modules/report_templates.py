"""
FEPD - Forensic Evidence Parser Dashboard
Report Template System

Generates professional forensic reports using configurable Jinja2 templates.

Features:
    - Multiple report types (executive, technical, compliance, incident response)
    - Jinja2 template engine with custom filters
    - PDF, HTML, DOCX export formats
    - Charts and visualizations embedding
    - Chain of custody documentation
    - Configurable sections and layouts
    - Branding customization (logo, colors, headers/footers)
    - Multi-language report generation

Report Types:
    1. Executive Summary - High-level overview for management
    2. Technical Deep-Dive - Detailed technical analysis for analysts
    3. Compliance Audit - Regulatory compliance reporting
    4. Incident Response - IR team operational report

Architecture:
    - ReportGenerator: Main interface
    - TemplateRenderer: Jinja2 rendering engine
    - ReportFormatter: Output format conversion (PDF, DOCX)
    - ChartEmbedder: Visualization integration
    - ReportConfig: Template configuration

Usage:
    from src.modules.report_templates import ReportGenerator
    
    # Initialize
    generator = ReportGenerator()
    
    # Generate executive summary
    report = generator.generate_report(
        report_type='executive',
        data=timeline_data,
        format='pdf',
        output_path='report.pdf'
    )

Copyright (c) 2025 FEPD Development Team
License: Proprietary - For Forensic Use Only
"""

import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field
from datetime import datetime
import json
import base64
from io import BytesIO

# Template engine
from jinja2 import Environment, FileSystemLoader, select_autoescape

# Document generation
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Visualization
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt


@dataclass
class ReportConfig:
    """
    Report configuration and branding.
    
    Attributes:
        title: Report title
        subtitle: Report subtitle
        organization: Organization name
        logo_path: Path to logo image
        analyst_name: Analyst name
        case_number: Case/incident number
        classification: Document classification (CONFIDENTIAL, etc.)
        colors: Brand color scheme
        include_toc: Include table of contents
        include_charts: Include visualization charts
        language: Report language code
    """
    title: str = "Forensic Analysis Report"
    subtitle: str = ""
    organization: str = "FEPD Forensics Team"
    logo_path: Optional[Path] = None
    analyst_name: str = "Forensic Analyst"
    case_number: str = ""
    classification: str = "CONFIDENTIAL"
    colors: Dict[str, str] = field(default_factory=lambda: {
        'primary': '#2c3e50',
        'secondary': '#3498db',
        'success': '#27ae60',
        'warning': '#f39c12',
        'danger': '#e74c3c'
    })
    include_toc: bool = True
    include_charts: bool = True
    language: str = 'en_US'


@dataclass
class ReportSection:
    """
    Report section with content and metadata.
    
    Attributes:
        title: Section title
        content: Section content (HTML or plain text)
        subsections: List of subsections
        charts: Embedded charts
        level: Heading level (1-6)
    """
    title: str
    content: str = ""
    subsections: List['ReportSection'] = field(default_factory=list)
    charts: List[str] = field(default_factory=list)  # Base64 encoded images
    level: int = 2


class ChartEmbedder:
    """
    Generates and embeds charts in reports.
    """
    
    def __init__(self, logger: Optional[logging.Logger] = None):
        self.logger = logger or logging.getLogger(__name__)
    
    def generate_timeline_chart(
        self,
        events: List[Dict[str, Any]],
        title: str = "Event Timeline"
    ) -> str:
        """
        Generate timeline chart.
        
        Args:
            events: List of events
            title: Chart title
        
        Returns:
            Base64 encoded PNG image
        """
        fig, ax = plt.subplots(figsize=(12, 6))
        
        # Extract data
        timestamps = [datetime.fromisoformat(e['timestamp']) if isinstance(e['timestamp'], str) else e['timestamp'] 
                     for e in events]
        categories = [e.get('category', 'Unknown') for e in events]
        
        # Color map
        category_colors = {
            'File Activity': '#3498db',
            'Network Activity': '#e74c3c',
            'Process Execution': '#f39c12',
            'Registry Modification': '#9b59b6',
            'Authentication': '#27ae60',
            'System Event': '#95a5a6'
        }
        
        colors = [category_colors.get(cat, '#95a5a6') for cat in categories]
        
        # Plot
        ax.scatter(timestamps, range(len(timestamps)), c=colors, s=100, alpha=0.6)
        ax.set_xlabel('Time')
        ax.set_ylabel('Event #')
        ax.set_title(title)
        ax.grid(True, alpha=0.3)
        
        # Legend
        unique_categories = list(set(categories))
        legend_elements = [plt.Line2D([0], [0], marker='o', color='w',  # type: ignore[attr-defined]
                                     markerfacecolor=category_colors.get(cat, '#95a5a6'), 
                                     markersize=10, label=cat)
                          for cat in unique_categories]
        ax.legend(handles=legend_elements, loc='best')
        
        # Encode to base64
        buffer = BytesIO()
        plt.tight_layout()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        
        return base64.b64encode(buffer.read()).decode('utf-8')
    
    def generate_category_distribution(
        self,
        events: List[Dict[str, Any]],
        title: str = "Event Distribution by Category"
    ) -> str:
        """Generate pie chart of event categories."""
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Count categories
        categories = {}
        for event in events:
            cat = event.get('category', 'Unknown')
            categories[cat] = categories.get(cat, 0) + 1
        
        # Plot
        colors = ['#3498db', '#e74c3c', '#f39c12', '#9b59b6', '#27ae60', '#95a5a6']
        ax.pie(categories.values(), labels=categories.keys(), autopct='%1.1f%%',
               colors=colors[:len(categories)], startangle=90)
        ax.set_title(title)
        
        # Encode
        buffer = BytesIO()
        plt.tight_layout()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        
        return base64.b64encode(buffer.read()).decode('utf-8')
    
    def generate_severity_distribution(
        self,
        events: List[Dict[str, Any]],
        title: str = "Event Distribution by Severity"
    ) -> str:
        """Generate bar chart of severity levels."""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # Count severities
        severities = {'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0, 'INFO': 0}
        for event in events:
            sev = event.get('severity', 'INFO')
            severities[sev] = severities.get(sev, 0) + 1
        
        # Plot
        colors = {'CRITICAL': '#e74c3c', 'HIGH': '#f39c12', 'MEDIUM': '#f1c40f', 
                 'LOW': '#3498db', 'INFO': '#95a5a6'}
        
        bars = ax.bar(severities.keys(), severities.values(),
                     color=[colors[s] for s in severities.keys()])
        
        ax.set_xlabel('Severity')
        ax.set_ylabel('Count')
        ax.set_title(title)
        ax.grid(True, alpha=0.3, axis='y')
        
        # Encode
        buffer = BytesIO()
        plt.tight_layout()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        
        return base64.b64encode(buffer.read()).decode('utf-8')


class TemplateRenderer:
    """
    Jinja2 template rendering engine.
    """
    
    def __init__(
        self,
        templates_dir: Path,
        logger: Optional[logging.Logger] = None
    ):
        """
        Initialize template renderer.
        
        Args:
            templates_dir: Directory containing Jinja2 templates
            logger: Optional logger
        """
        self.logger = logger or logging.getLogger(__name__)
        self.templates_dir = templates_dir
        
        # Initialize Jinja2 environment
        self.env = Environment(
            loader=FileSystemLoader(str(templates_dir)),
            autoescape=select_autoescape(['html', 'xml'])
        )
        
        # Add custom filters
        self.env.filters['datetime'] = self._format_datetime
        self.env.filters['filesize'] = self._format_filesize
        self.env.filters['duration'] = self._format_duration
        self.env.filters['severity_color'] = self._get_severity_color
        
        self.logger.info(f"Template renderer initialized: {templates_dir}")
    
    def _format_datetime(self, dt: datetime, format: str = '%Y-%m-%d %H:%M:%S') -> str:
        """Format datetime filter."""
        if isinstance(dt, str):
            dt = datetime.fromisoformat(dt)
        return dt.strftime(format)
    
    def _format_filesize(self, size_bytes: int) -> str:
        """Format file size filter."""
        size: float = float(size_bytes)
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size < 1024.0:
                return f"{size:.2f} {unit}"
            size /= 1024.0
        return f"{size:.2f} PB"
    
    def _format_duration(self, seconds: float) -> str:
        """Format duration filter."""
        if seconds < 60:
            return f"{seconds:.1f} sec"
        elif seconds < 3600:
            return f"{seconds/60:.1f} min"
        elif seconds < 86400:
            return f"{seconds/3600:.1f} hours"
        else:
            return f"{seconds/86400:.1f} days"
    
    def _get_severity_color(self, severity: str) -> str:
        """Get color for severity level."""
        colors = {
            'CRITICAL': '#e74c3c',
            'HIGH': '#f39c12',
            'MEDIUM': '#f1c40f',
            'LOW': '#3498db',
            'INFO': '#95a5a6'
        }
        return colors.get(severity, '#95a5a6')
    
    def render(
        self,
        template_name: str,
        context: Dict[str, Any]
    ) -> str:
        """
        Render template with context.
        
        Args:
            template_name: Template filename
            context: Template variables
        
        Returns:
            Rendered HTML string
        """
        try:
            template = self.env.get_template(template_name)
            return template.render(**context)
        
        except Exception as e:
            self.logger.error(f"Template rendering failed: {e}")
            raise


class ReportGenerator:
    """
    Main report generation interface.
    """
    
    def __init__(
        self,
        templates_dir: Optional[Path] = None,
        logger: Optional[logging.Logger] = None
    ):
        """
        Initialize report generator.
        
        Args:
            templates_dir: Directory containing templates
            logger: Optional logger
        """
        self.logger = logger or logging.getLogger(__name__)
        
        # Set templates directory
        if templates_dir is None:
            templates_dir = Path(__file__).parent.parent.parent / 'templates'
        self.templates_dir = Path(templates_dir)
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize components
        self.renderer = TemplateRenderer(self.templates_dir, logger=logger)
        self.chart_embedder = ChartEmbedder(logger=logger)
        
        self.logger.info("Report generator initialized")
    
    def generate_report(
        self,
        report_type: str,
        data: Dict[str, Any],
        config: Optional[ReportConfig] = None,
        format: str = 'html',
        output_path: Optional[Path] = None
    ) -> str:
        """
        Generate forensic report.
        
        Args:
            report_type: Type of report ('executive', 'technical', 'compliance', 'incident')
            data: Report data (events, analysis results, etc.)
            config: Report configuration
            format: Output format ('html', 'pdf', 'docx')
            output_path: Output file path (optional)
        
        Returns:
            Generated report content (HTML or path to file)
        
        Example:
            generator = ReportGenerator()
            
            # Generate executive summary PDF
            report = generator.generate_report(
                report_type='executive',
                data={'events': timeline_events, 'analysis': ml_results},
                format='pdf',
                output_path=Path('executive_summary.pdf')
            )
        """
        # Use default config if not provided
        if config is None:
            config = ReportConfig()
        
        # Validate report type
        valid_types = ['executive', 'technical', 'compliance', 'incident']
        if report_type not in valid_types:
            raise ValueError(f"Invalid report type: {report_type}. Must be one of {valid_types}")
        
        # Prepare context
        context = self._prepare_context(report_type, data, config)
        
        # Render template
        template_name = f"{report_type}_report.html"
        html_content = self.renderer.render(template_name, context)
        
        # Convert to requested format
        if format == 'html':
            if output_path:
                output_path.parent.mkdir(parents=True, exist_ok=True)
                output_path.write_text(html_content, encoding='utf-8')
                self.logger.info(f"Generated HTML report: {output_path}")
                return str(output_path)
            else:
                return html_content
        
        elif format == 'pdf':
            return self._convert_to_pdf(html_content, output_path)
        
        elif format == 'docx':
            return self._convert_to_docx(context, output_path)
        
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _prepare_context(
        self,
        report_type: str,
        data: Dict[str, Any],
        config: ReportConfig
    ) -> Dict[str, Any]:
        """Prepare template context."""
        events = data.get('events', [])
        
        # Base context
        context = {
            'config': config,
            'report_type': report_type,
            'generated_at': datetime.now(),
            'total_events': len(events),
            'events': events,
            'data': data
        }
        
        # Add charts if enabled
        if config.include_charts and events:
            context['charts'] = {
                'timeline': self.chart_embedder.generate_timeline_chart(events),
                'category_dist': self.chart_embedder.generate_category_distribution(events),
                'severity_dist': self.chart_embedder.generate_severity_distribution(events)
            }
        
        # Add report-specific context
        if report_type == 'executive':
            context['summary'] = self._generate_executive_summary(events, data)
        
        elif report_type == 'technical':
            context['analysis'] = self._generate_technical_analysis(events, data)
        
        elif report_type == 'compliance':
            context['compliance'] = self._generate_compliance_section(events, data)
        
        elif report_type == 'incident':
            context['incident'] = self._generate_incident_details(events, data)
        
        return context
    
    def _generate_executive_summary(
        self,
        events: List[Dict[str, Any]],
        data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate executive summary section."""
        # Count by severity
        severity_counts = {'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0, 'INFO': 0}
        for event in events:
            sev = event.get('severity', 'INFO')
            severity_counts[sev] = severity_counts.get(sev, 0) + 1
        
        # Key findings
        key_findings = []
        if severity_counts['CRITICAL'] > 0:
            key_findings.append(f"{severity_counts['CRITICAL']} CRITICAL security events detected")
        if severity_counts['HIGH'] > 0:
            key_findings.append(f"{severity_counts['HIGH']} HIGH severity incidents identified")
        
        # Anomalies
        anomalies = data.get('anomalies', [])
        if anomalies:
            key_findings.append(f"{len(anomalies)} behavioral anomalies detected")
        
        # Threat matches
        threat_matches = data.get('threat_matches', [])
        if threat_matches:
            key_findings.append(f"{len(threat_matches)} threat intelligence matches found")
        
        return {
            'severity_counts': severity_counts,
            'key_findings': key_findings,
            'risk_score': self._calculate_risk_score(severity_counts),
            'recommendations': self._generate_recommendations(severity_counts, anomalies, threat_matches)
        }
    
    def _generate_technical_analysis(
        self,
        events: List[Dict[str, Any]],
        data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate technical analysis section."""
        return {
            'event_breakdown': self._analyze_events_by_category(events),
            'timeline_analysis': self._analyze_timeline(events),
            'anomaly_details': data.get('anomaly_details', []),
            'ueba_alerts': data.get('ueba_alerts', []),
            'threat_intelligence': data.get('threat_matches', []),
            'artifacts_analyzed': data.get('artifacts', [])
        }
    
    def _generate_compliance_section(
        self,
        events: List[Dict[str, Any]],
        data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate compliance section."""
        return {
            'frameworks': ['PCI-DSS', 'HIPAA', 'GDPR', 'SOX'],
            'audit_trail': events,
            'violations': data.get('violations', []),
            'chain_of_custody': data.get('chain_of_custody', [])
        }
    
    def _generate_incident_details(
        self,
        events: List[Dict[str, Any]],
        data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate incident response details."""
        return {
            'incident_timeline': events,
            'affected_systems': data.get('affected_systems', []),
            'indicators_of_compromise': data.get('iocs', []),
            'response_actions': data.get('actions', []),
            'containment_status': data.get('containment', 'In Progress')
        }
    
    def _analyze_events_by_category(self, events: List[Dict[str, Any]]) -> Dict[str, int]:
        """Analyze events by category."""
        categories = {}
        for event in events:
            cat = event.get('category', 'Unknown')
            categories[cat] = categories.get(cat, 0) + 1
        return categories
    
    def _analyze_timeline(self, events: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze timeline characteristics."""
        if not events:
            return {}
        
        timestamps = [datetime.fromisoformat(e['timestamp']) if isinstance(e['timestamp'], str) else e['timestamp'] 
                     for e in events]
        
        return {
            'earliest': min(timestamps),
            'latest': max(timestamps),
            'duration': (max(timestamps) - min(timestamps)).total_seconds(),
            'events_per_hour': len(events) / max(1, (max(timestamps) - min(timestamps)).total_seconds() / 3600)
        }
    
    def _calculate_risk_score(self, severity_counts: Dict[str, int]) -> int:
        """Calculate overall risk score (0-100)."""
        weights = {'CRITICAL': 50, 'HIGH': 20, 'MEDIUM': 5, 'LOW': 1, 'INFO': 0}
        score = sum(count * weights.get(sev, 0) for sev, count in severity_counts.items())
        return min(100, score)
    
    def _generate_recommendations(
        self,
        severity_counts: Dict[str, int],
        anomalies: List[Any],
        threat_matches: List[Any]
    ) -> List[str]:
        """Generate recommendations."""
        recs = []
        
        if severity_counts['CRITICAL'] > 0:
            recs.append("Immediate incident response required for critical events")
        
        if anomalies:
            recs.append("Review anomalous behaviors for potential insider threats")
        
        if threat_matches:
            recs.append("Isolate systems matching threat intelligence indicators")
        
        recs.append("Continue monitoring for additional suspicious activity")
        recs.append("Update security controls based on findings")
        
        return recs
    
    def _convert_to_pdf(self, html_content: str, output_path: Optional[Path]) -> str:
        """Convert HTML to PDF using WeasyPrint."""
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("WeasyPrint not available. Install with: pip install weasyprint")
        
        if output_path is None:
            output_path = Path(f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Convert
        HTML(string=html_content).write_pdf(output_path)
        
        self.logger.info(f"Generated PDF report: {output_path}")
        return str(output_path)
    
    def _convert_to_docx(self, context: Dict[str, Any], output_path: Optional[Path]) -> str:
        """Convert to DOCX format."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        if output_path is None:
            output_path = Path(f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create document
        doc = Document()
        
        # Add title
        config = context['config']
        title = doc.add_heading(config.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        if config.subtitle:
            subtitle = doc.add_paragraph(config.subtitle)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated: {context['generated_at'].strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Analyst: {config.analyst_name}")
        if config.case_number:
            doc.add_paragraph(f"Case #: {config.case_number}")
        
        doc.add_page_break()
        
        # Add content sections
        doc.add_heading('Executive Summary', level=1)
        if 'summary' in context:
            summary = context['summary']
            doc.add_paragraph(f"Total Events: {context['total_events']}")
            doc.add_paragraph(f"Risk Score: {summary.get('risk_score', 0)}/100")
        
        # Save
        doc.save(output_path)
        
        self.logger.info(f"Generated DOCX report: {output_path}")
        return str(output_path)


# Example usage
if __name__ == '__main__':
    # Example: Generate executive summary
    generator = ReportGenerator()
    
    # Mock data
    mock_events = [
        {
            'id': 'evt_001',
            'timestamp': '2025-11-07 10:30:00',
            'category': 'File Activity',
            'severity': 'HIGH',
            'description': 'Sensitive file accessed'
        },
        {
            'id': 'evt_002',
            'timestamp': '2025-11-07 10:31:00',
            'category': 'Network Activity',
            'severity': 'CRITICAL',
            'description': 'Data exfiltration detected'
        }
    ]
    
    # Generate report
    report = generator.generate_report(
        report_type='executive',
        data={'events': mock_events},
        format='html',
        output_path=Path('executive_summary.html')
    )
    
    print(f"Report generated: {report}")
